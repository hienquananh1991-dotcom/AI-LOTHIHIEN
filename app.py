import streamlit as st
import time
import random
from docx import Document
from io import BytesIO
from PIL import Image

# --- 1. CẤU HÌNH & DỮ LIỆU ---
st.set_page_config(page_title="Smart-Print AI Điện Biên", page_icon="🏫", layout="wide")

# Dữ liệu giả lập Mục tiêu bài học (SGK Kết nối tri thức)
MUC_TIEU_SGK = {
    "Toán": "Thực hiện được phép cộng, trừ, nhân, chia. Giải quyết được vấn đề gắn với thực tiễn.",
    "Tiếng Việt": "Đọc trôi chảy, hiểu nội dung văn bản. Viết đúng chính tả và ngữ pháp.",
    "Tin học": "Bước đầu làm quen với thiết bị số. Biết bảo vệ sức khỏe khi sử dụng máy tính.",
    "Công nghệ": "Sử dụng được vật liệu thủ công. Nhận biết được một số sản phẩm công nghệ.",
    "Khoa học": "Khám phá thế giới tự nhiên. Biết cách chăm sóc sức khỏe bản thân.",
    "Lịch sử & Địa lý": "Nhận biết được cảnh quan thiên nhiên và di tích lịch sử địa phương.",
    "Tiếng Anh": "Nghe, nói, đọc, viết các từ vựng và mẫu câu cơ bản theo chủ đề.",
    "Đạo đức": "Biết yêu thương gia đình, thầy cô, bạn bè. Trung thực trong học tập.",
    "Mĩ thuật": "Biết sử dụng màu sắc, đường nét để tạo hình sản phẩm đơn giản.",
    "Âm nhạc": "Hát đúng giai điệu, lời ca. Biết vận động theo nhịp điệu bài hát.",
    "Thể dục": "Thực hiện được các động tác đội hình đội ngũ và bài tập rèn luyện tư thế."
}

# Hàm tạo file Word
def tao_file_word(ten_hs, mon_hoc, lop, noi_dung, loi_khuyen, muc_tieu):
    doc = Document()
    doc.add_heading(f'PHIẾU BÀI TẬP: {ten_hs.upper()}', 0)
    doc.add_paragraph(f'Môn: {mon_hoc} - {lop}')
    doc.add_paragraph('Bộ sách: Kết nối tri thức với cuộc sống')
    doc.add_paragraph(f'Mục tiêu bài học: {muc_tieu}')
    doc.add_paragraph('-'*50)
    doc.add_heading('A. BÀI TẬP THỰC HÀNH', level=1)
    doc.add_paragraph(noi_dung)
    doc.add_heading('B. GÓC SƯ PHẠM (AI Gợi ý)', level=1)
    doc.add_paragraph(f"Lời khuyên: {loi_khuyen}")
    doc.add_paragraph('\n')
    doc.add_paragraph('--- Smart-Print AI: Đồng hành cùng giáo dục vùng cao ---')
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 2. GIAO DIỆN CHÍNH ---
st.title("🏫 Smart-Print AI: Hệ Sinh Thái Giáo Dục Số")
st.markdown("**Địa phương:** Tỉnh Điện Biên | **Bộ sách:** Kết nối tri thức với cuộc sống")
st.markdown("---")

# Sidebar nhập liệu chung
with st.sidebar:
    st.image("https://upload.wikimedia.org/wikipedia/commons/thumb/1/11/FPT_logo_2010.svg/1200px-FPT_logo_2010.svg.png", width=100, caption="Logo Trường/Dự án")
    st.header("👤 Hồ sơ học sinh")
    ten_hs = st.text_input("Họ tên", "Lò Thị Mai")
    lop = st.selectbox("Khối lớp", ["Lớp 1", "Lớp 2", "Lớp 3", "Lớp 4", "Lớp 5"])
    hoc_luc = st.select_slider("Mức độ năng lực", options=["Yếu", "Trung bình", "Khá", "Giỏi"])
    st.info("💡 Hệ thống tự động liên kết dữ liệu với kho học liệu số.")

# TẠO 3 TAB CHỨC NĂNG
tab1, tab2, tab3 = st.tabs(["📝 SOẠN BÀI & TẢI VỀ", "📷 CHẤM BÀI QUA ẢNH (AI)", "📚 TRA CỨU SGK"])

# --- TAB 1: SOẠN BÀI ---
with tab1:
    col_mon, col_action = st.columns([3, 1])
    with col_mon:
        mon_hoc = st.selectbox("Chọn môn học:", 
            ["Toán", "Tiếng Việt", "Tin học", "Công nghệ", "Khoa học", 
             "Lịch sử & Địa lý", "Tiếng Anh", "Đạo đức", "Mĩ thuật", "Âm nhạc", "Thể dục"])
    with col_action:
        st.write("") # Spacer
        st.write("")
        btn_tao = st.button("🚀 TẠO PHIẾU", type="primary", use_container_width=True)

    if btn_tao:
        with st.spinner(f"Đang tham chiếu SGK {mon_hoc} để soạn bài..."):
            time.sleep(1.5)
            # Demo nội dung
            muc_tieu = MUC_TIEU_SGK.get(mon_hoc, "Bám sát chương trình GDPT 2018")
            
            if mon_hoc == "Toán":
                noi_dung = "Bài 1: Tính nhẩm...\nBài 2: Giải toán có lời văn về thu hoạch nông sản..."
            elif mon_hoc == "Tin học":
                noi_dung = "Câu 1: Em hãy khoanh tròn vào thiết bị là máy tính.\nCâu 2: Tư thế ngồi đúng..."
            elif mon_hoc == "Thể dục":
                noi_dung = "Hoạt động: Thực hiện động tác vươn thở và tay (Mỗi động tác 2 lần 8 nhịp)."
            else:
                noi_dung = f"Câu hỏi ôn tập kiến thức môn {mon_hoc} tuần này.\nHoạt động thực hành tại nhà/bản làng."
            
            loi_khuyen = "Hãy khen ngợi khi em hoàn thành nhiệm vụ."

            # Hiển thị
            st.success("✅ Đã tạo xong!")
            with st.expander("👀 Xem trước nội dung phiếu"):
                st.write(f"**Mục tiêu:** {muc_tieu}")
                st.code(noi_dung, language=None)
            
            # Tải về
            file_word = tao_file_word(ten_hs, mon_hoc, lop, noi_dung, loi_khuyen, muc_tieu)
            st.download_button("📥 TẢI PHIẾU WORD (.docx)", file_word, f"{ten_hs}_{mon_hoc}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# --- TAB 2: CHẤM BÀI QUA ẢNH (AI VISION) ---
with tab2:
    st.header("🤖 Trợ lý chấm bài & Nhận xét")
    st.write("Tải lên hình ảnh phiếu bài tập học sinh đã làm (chụp bằng điện thoại).")
    
    uploaded_file = st.file_uploader("Chọn ảnh bài làm...", type=['jpg', 'png', 'jpeg'])
    
    if uploaded_file is not None:
        col_img, col_result = st.columns(2)
        
        with col_img:
            image = Image.open(uploaded_file)
            st.image(image, caption='Bài làm của học sinh', use_column_width=True)
            btn_cham = st.button("✨ AI PHÂN TÍCH & CHẤM ĐIỂM")
            
        with col_result:
            if btn_cham:
                with st.spinner("AI đang đọc chữ viết tay và so sánh đáp án..."):
                    time.sleep(2) # Giả lập thời gian xử lý
                    
                    # KẾT QUẢ GIẢ LẬP (MÔ PHỎNG)
                    diem = random.randint(6, 10)
                    nhan_xet = ""
                    if diem >= 9:
                        nhan_xet = "Em làm bài rất tốt! Chữ viết sạch đẹp. Đã hiểu rõ mục tiêu bài học."
                        color = "green"
                    elif diem >= 7:
                        nhan_xet = "Em làm bài khá. Tuy nhiên cần chú ý lỗi chính tả ở câu 2."
                        color = "orange"
                    else:
                        nhan_xet = "Em cần cố gắng hơn. Chưa nắm vững kiến thức cơ bản."
                        color = "red"
                    
                    st.markdown(f"### Kết quả: :{color}[{diem}/10 điểm]")
                    st.info(f"**Nhận xét chi tiết:**\n{nhan_xet}")
                    
                    st.markdown("**Đánh giá mức độ đạt mục tiêu:**")
                    st.progress(diem * 10)
                    st.caption(f"Căn cứ theo chuẩn kiến thức kĩ năng môn {mon_hoc}.")

# --- TAB 3: TRA CỨU SGK (LIÊN KẾT) ---
with tab3:
    st.header("📖 Kết nối tri thức với cuộc sống")
    st.write("Hệ thống tự động trích xuất mục tiêu bài học để giáo viên đối chiếu.")
    
    col_sgk_1, col_sgk_2 = st.columns([2, 1])
    
    with col_sgk_1:
        st.subheader(f"Mục tiêu môn: {mon_hoc}")
        st.success(MUC_TIEU_SGK.get(mon_hoc, "Đang cập nhật dữ liệu..."))
        
        st.markdown("### Gợi ý phương pháp dạy học:")
        st.markdown("- **Phương pháp trực quan:** Sử dụng tranh ảnh, vật thật (ngô, khoai, sắn...).")
        st.markdown("- **Phương pháp trò chơi:** 'Rung chuông vàng', 'Ai nhanh hơn'.")
        
    with col_sgk_2:
        st.info("🔗 **Nguồn tài liệu chính thống**")
        st.write("Để xem chi tiết từng trang sách, thầy cô vui lòng truy cập Hành trang số (NXB Giáo dục):")
        st.link_button("🌐 Truy cập Hành Trang Số", "https://hanhtrangso.nxbgd.vn/")
        st.image("https://hanhtrangso.nxbgd.vn/img/logo.png", width=150)

# --- FOOTER ---
st.markdown("---")
st.caption("© 2024 Dự án Chuyển đổi số Giáo dục Điện Biên. Powered by Streamlit & AI.")
